import os
import json
from datetime import datetime
import numpy as np
import pandas as pd
import yfinance as yf
import matplotlib.pyplot as plt

from backtesting import Backtest, Strategy

import re
from docx import Document
from openai import AzureOpenAI

from dotenv import load_dotenv
load_dotenv()


# -----------------------
# Config
# -----------------------
START_CASH = 10_000
COMMISSION = 0.001  # 0.1%

REQUIRED_COLS = ["50_sma", "200_sma", "RSI_14", "MACD_Hist", "ADX_14", "Close", "regime","Adj Close"]


# -----------------------
# Data + Indicators
# -----------------------
import time

def download_ohlcv(
    ticker: str,
    buffer_start: str,
    eval_start: str,
    eval_end: str,
    max_retries: int = 3,
    pause_sec: float = 2.0,
) -> pd.DataFrame:
    last_err = None

    for attempt in range(1, max_retries + 1):
        try:
            df = yf.download(
                ticker,
                start=buffer_start,
                end=eval_end,
                auto_adjust=False,
                progress=False,
                threads=False,
            )

            if df is None or df.empty:
                raise ConnectionError(
                    f"yfinance returned empty data for {ticker} "
                    f"(attempt {attempt}/{max_retries})"
                )

            if isinstance(df.columns, pd.MultiIndex):
                df.columns = df.columns.get_level_values(0)

            needed = {"Open", "High", "Low", "Close", "Volume", "Adj Close"}
            missing = needed - set(df.columns)
            if missing:
                raise ValueError(f"Missing columns from yfinance: {missing}")

            df = df[["Open", "High", "Low", "Close", "Volume", "Adj Close"]].copy()
            df.index = pd.to_datetime(df.index)

            return df

        except Exception as e:
            last_err = e
            if attempt < max_retries:
                time.sleep(pause_sec * attempt)

    raise RuntimeError(
        f"Failed to download {ticker} after {max_retries} attempts. "
        f"Last error: {last_err}"
    )


def compute_indicators(df: pd.DataFrame) -> pd.DataFrame:
    # SMA + regime
    df["50_sma"] = df["Close"].rolling(50).mean()
    df["200_sma"] = df["Close"].rolling(200).mean()
    df["regime"] = np.where(df["50_sma"] > df["200_sma"], 1, 0)

    # --- RSI(14) ---
    delta = df["Close"].diff()
    gain = delta.clip(lower=0)
    loss = (-delta).clip(lower=0)
    avg_gain = gain.rolling(14).mean()
    avg_loss = loss.rolling(14).mean()
    rs = avg_gain / avg_loss.replace(0, np.nan)
    df["RSI_14"] = 100 - (100 / (1 + rs))

    # --- MACD histogram (12,26,9) ---
    ema12 = df["Close"].ewm(span=12, adjust=False).mean()
    ema26 = df["Close"].ewm(span=26, adjust=False).mean()
    macd = ema12 - ema26
    signal = macd.ewm(span=9, adjust=False).mean()
    df["MACD_Hist"] = macd - signal

    # --- ADX(14) ---
    df["tr1"] = df["High"] - df["Low"]
    df["tr2"] = (df["High"] - df["Close"].shift(1)).abs()
    df["tr3"] = (df["Low"] - df["Close"].shift(1)).abs()
    df["TR"] = df[["tr1", "tr2", "tr3"]].max(axis=1)

    # Directional Movement (DM)
    df["up_move"] = df["High"] - df["High"].shift(1)
    df["down_move"] = df["Low"].shift(1) - df["Low"]

    df["+DM"] = np.where(
        (df["up_move"] > df["down_move"]) & (df["up_move"] > 0),
        df["up_move"],
        0,
    )
    df["-DM"] = np.where(
        (df["down_move"] > df["up_move"]) & (df["down_move"] > 0),
        df["down_move"],
        0,
    )

    period = 14
    df["TR_smooth"] = df["TR"].ewm(alpha=1 / period, adjust=False).mean()
    df["+DM_smooth"] = df["+DM"].ewm(alpha=1 / period, adjust=False).mean()
    df["-DM_smooth"] = df["-DM"].ewm(alpha=1 / period, adjust=False).mean()

    df["+DI"] = 100 * (df["+DM_smooth"] / df["TR_smooth"])
    df["-DI"] = 100 * (df["-DM_smooth"] / df["TR_smooth"])

    # ADX
    df["DX"] = 100 * ((df["+DI"] - df["-DI"]).abs() / (df["+DI"] + df["-DI"]))
    df["ADX_14"] = df["DX"].ewm(alpha=1 / period, adjust=False).mean()

    return df


# -----------------------
# Signals + Position
# -----------------------
def compute_signals_and_position(df: pd.DataFrame) -> pd.DataFrame:
    trend_entry = (
        (df["regime"] == 1) &
        (df["ADX_14"] > 19) & (df["ADX_14"] < 30) &
        (df["RSI_14"] > 35) & (df["RSI_14"] < 69) &
        (df["MACD_Hist"] > -0.025)
    )

    trend_exit = (
        (df["regime"] != 1) |
        ((df["MACD_Hist"] < -0.025) & (df["ADX_14"] <= 19)) |
        (df["RSI_14"] > 70)
    )

    df["entry"] = trend_entry
    df["exit"] = trend_exit

    pos = np.zeros(len(df), dtype=float)
    in_pos = 0.0
    for i in range(len(df)):
        if in_pos == 0.0 and df["entry"].iloc[i]:
            in_pos = 1.0
        elif in_pos == 1.0 and df["exit"].iloc[i]:
            in_pos = 0.0
        pos[i] = in_pos

    df["position"] = pos
    df["position_shifted"] = df["position"].shift(1).fillna(0)

    df["recommendation"] = "HOLD"
    df.loc[(df["entry"]) & (df["position_shifted"] == 0), "recommendation"] = "BUY"
    df.loc[(df["exit"]) & (df["position_shifted"] == 1), "recommendation"] = "SELL"

    return df


# -----------------------
# Strategy equity (Adj Close, net commission, cash-based)
# -----------------------
def compute_strategy_equity_adj_net(
    df: pd.DataFrame,
    start_cash: float,
    commission: float,
    required_cols: list[str],
    eval_start: str,
    eval_end: str,
):
    # 1) Build evaluation dataframe 
    df_eval = df.dropna(subset=required_cols).loc[eval_start:eval_end].copy()
    if df_eval.empty:
        raise RuntimeError("df_eval is empty after dropna() + date slicing. Check eval_start/eval_end and required_cols.")

    # 2) Adj Close returns (eval window)
    df_eval["adj_ret"] = df_eval["Adj Close"].pct_change().fillna(0)

    # 3) Align position within eval window 
    df_eval["position_shifted"] = df_eval["position_shifted"].reindex(df_eval.index).ffill().fillna(0)

    # 4) Returns + commission
    df_eval["strategy_adj_ret_gross"] = df_eval["position_shifted"] * df_eval["adj_ret"]
    df_eval["turnover"] = df_eval["position_shifted"].diff().fillna(0).abs()
    df_eval["strategy_adj_ret_net"] = df_eval["strategy_adj_ret_gross"] - df_eval["turnover"] * commission

    # 5) Equity curve 
    df_eval["strategy_adj_equity_net"] = start_cash * (1 + df_eval["strategy_adj_ret_net"]).cumprod()

    # 6) Metrics
    total_return_pct = (df_eval["strategy_adj_equity_net"].iloc[-1] / start_cash - 1) * 100

    days = (df_eval.index[-1] - df_eval.index[0]).days
    years = days / 365.25
    cagr_pct = ((df_eval["strategy_adj_equity_net"].iloc[-1] / start_cash) ** (1 / years) - 1) * 100 if years > 0 else float("nan")

    strat_ret = df_eval["strategy_adj_ret_net"]
    sharpe = (strat_ret.mean() / strat_ret.std()) * np.sqrt(252) if strat_ret.std() != 0 else float("nan")

    strat_equity = df_eval["strategy_adj_equity_net"]
    strat_roll_max = strat_equity.cummax()
    strat_drawdown = (strat_equity / strat_roll_max) - 1
    max_dd_pct = strat_drawdown.min() * 100

    return df_eval, {
        "total_return_pct": float(total_return_pct),
        "cagr_pct": float(cagr_pct),
        "sharpe": float(sharpe),
        "max_drawdown_pct": float(max_dd_pct),
    }


def compute_buy_hold_adj_net(
    df: pd.DataFrame,
    start_cash: float,
    commission: float,
    required_cols: list,
    eval_start: str,
    eval_end: str,
):
    # 1) Build evaluation frame
    df_eval = df.dropna(subset=required_cols).loc[eval_start:eval_end].copy()

    if df_eval.empty:
        raise ValueError("df_eval is empty after dropna + slicing eval_start:eval_end")

    # 2) Adj Close daily returns
    df_eval["adj_ret"] = df_eval["Adj Close"].pct_change().fillna(0)

    # 3) Buy & Hold index
    bh_idx_adj_close = df_eval["Adj Close"] / df_eval["Adj Close"].iloc[0]

    # 4) Entry-only commission
    bh_idx_adj_close_net = bh_idx_adj_close * (1 - commission)

    # 5) Equity in cash
    bh_equity = start_cash * bh_idx_adj_close_net

    # 6) Metrics
    bh_total_return_pct = (bh_equity.iloc[-1] / start_cash - 1) * 100

    days_held = (bh_equity.index[-1] - bh_equity.index[0]).days
    years_held = days_held / 365.25
    bh_cagr_pct = ((bh_equity.iloc[-1] / start_cash) ** (1 / years_held) - 1) * 100 if years_held > 0 else np.nan

    bh_ret = bh_equity.pct_change().fillna(0)
    bh_sharpe = (bh_ret.mean() / bh_ret.std()) * np.sqrt(252) if bh_ret.std() != 0 else float("nan")

    bh_roll_max = bh_equity.cummax()
    bh_drawdown = (bh_equity / bh_roll_max) - 1
    bh_max_drawdown_pct = bh_drawdown.min() * 100

    return bh_equity, {
        "total_return_pct": float(bh_total_return_pct),
        "cagr_pct": float(bh_cagr_pct),
        "sharpe": float(bh_sharpe),
        "max_drawdown_pct": float(bh_max_drawdown_pct),
    }

# ---------------------------------------------------------------
# backtesting.py Strategy (Close-to-close performance evaluation)
# ---------------------------------------------------------------
class KOstrategy(Strategy):
    def init(self):
        pass

    def next(self):
        if not self.position:
            if self.data["entry"][-1]:
                self.buy()
        else:
            if self.data["exit"][-1]:
                self.position.close()


def run_backtesting_py(df_bt: pd.DataFrame, cash: float, commission: float):
    bt = Backtest(df_bt, KOstrategy, cash=cash, commission=commission)
    stats = bt.run()
    trades = stats["_trades"].copy() if "_trades" in stats else pd.DataFrame()
    return stats, trades


# -----------------------
# Prompt (LLM)
# -----------------------
def build_user_prompt(df: pd.DataFrame, ticker: str, strat_metrics: dict, bh_metrics: dict, hit_rate_pct: float, n_trades: int) -> str:
    last = df.iloc[-1]
    regime = "TREND" if last["regime"] == 1 else "NO TREND"
    today = df.index[-1]
    today_recommendation = df.loc[today, "recommendation"]

    return f"""
Context: Analysis of the trading strategy for {ticker}.
Write a structured, professional two-page trade note based strictly on the information below.

Ticker: {ticker}
Date: {df.index[-1].date()}

Market Regime: {regime}
Recommendation: {today_recommendation}

Indicators:
- SMA50: {last["50_sma"]:.2f}
- SMA200: {last["200_sma"]:.2f}
- RSI(14): {last["RSI_14"]:.2f}
- MACD Histogram: {last["MACD_Hist"]:.2f}
- ADX: {last["ADX_14"]:.2f}

Performance Metrics (Net of Commission, Adj Close, Start Cash = {START_CASH:,}):

Strategy:
- Total Return (%): {strat_metrics["total_return_pct"]:.2f}
- CAGR (%): {strat_metrics["cagr_pct"]:.2f}
- Sharpe Ratio: {strat_metrics["sharpe"]:.2f}
- Max Drawdown (%): {strat_metrics["max_drawdown_pct"]:.2f}
- Hit Rate (%): {hit_rate_pct:.2f}
- Number of Trades: {n_trades}

Buy & Hold:
- Total Return (%): {bh_metrics["total_return_pct"]:.2f}
- CAGR (%): {bh_metrics["cagr_pct"]:.2f}
- Sharpe Ratio: {bh_metrics["sharpe"]:.2f}
- Max Drawdown (%): {bh_metrics["max_drawdown_pct"]:.2f}

Instructions:
- Do not invent numbers.
- Base all analysis strictly on the metrics and indicators above.
- Discuss relative performance versus Buy & Hold.
- Comment on risk-adjusted returns, drawdowns, and regime suitability.
- Maintain an academic yet practitioner-oriented tone.

This is my strategy: trend_entry when there is trend regime (i.e., when SMA50 > SMA200) , ADX is between 19 and 30, MACD Histogram is greater than -0.025 and when RSI (14) is not less than 35 and less than 69. On the other hand, the trend_exit when the trend is break (i.e., when SMA 50 <= SMA 200) or when ADX is not greater than 19 and MACD Histogram is below -0.025, or when RSI (14) is larger than 70.

1) Market regime assessment
2) Explanation of indicator confluence
3) Current trade recommendation (BUY / SELL / HOLD) based on my strategy
4) Backtest performance summary (CAGR, Sharpe, drawdown, hit rate)
5) Key risks and limitations

OUTPUT RULES:
- Return ONLY a valid JSON object (no markdown, no ``` fences).
- Use clear, report-style writing inside the fields:
  - Use short paragraphs.
  - Use bullet lists with leading "- " when appropriate.
  - Use concise section headings ONLY as plain text inside strings if needed (no markdown).

STYLE:
- Professional research note tone.
- Use subheadings within each section where helpful (e.g., "What we see:", "Why it matters:", "Decision:").
- Keep numbers exactly as provided. Do NOT invent data.


EXPECTED JSON STRUCTURE:
{{
  "Abstract": {{
    "Summary": "Max 100 words. State the analysis date and summarise recent performance using the latest indicators and signals."
  }},
  "Recommendation": {{
    "Signal": "Buy | Sell | Hold",
    "Rationale": "Detail justification. If there is a trend, but it is recommended to sell, please explain the exit logic clearly. Detail why one of the criteria can trigger for exit to ensure that it is not a early exit."
  }},
  "Strategy": {{
    "Overview": "Max 300 words. Explain the strategy logic clearly by separating entry rules, exit rules, and how the current market state maps to the recommendation.",
    "Indicators_used": "Describe SMA, RSI, MACD, ADX, and regime logic. First explain what each indicator measures in general, then interpret the current values and what they imply for current signal."
  }},
  "Backtest_summary": {{
    "Key_metrics": "CAGR, Sharpe ratio, maximum drawdown, hit rate.",
    "Performance_interpretation": " MAx 200 words interpreting results relative to Buy & Hold. Also, please mention the backtesting period i.e. the evaluation period. Detail 100 worss in explaning why the agent download more than 10 years raw data and then evaluate for exactly 10 years."
  }},
  "Limitations_and_risks": {{
    "Discussion": "Max 200 words discussing weaknesses, risks, and portfolio insights. Explain why the backtest is conducted using both Close and Adjusted Close prices, and justify what is the risk of the Adjusted Close results that are used for evaluation."
  }},
  "Future_refinements": {{
    "Discussion": "Max 200 words outlining future improvements, for example: walk-forward validation, and multiple indicators."
  }}
}}
""".strip()

def build_azure_client():
    return AzureOpenAI(
        api_version=os.environ["AZURE_API_VERSION"],
        azure_endpoint=os.environ["AZURE_ENDPOINT"],
        api_key=os.environ["AZURE_API_KEY"],
    )

def generate_trade_note_llm(client, deployment: str, system_prompt: str, user_prompt: str,
                            temperature: float = 0.2, max_tokens: int = 800) -> str:
    """Calls Azure OpenAI chat.completions and returns raw text."""
    resp = client.chat.completions.create(
        model=deployment,
        temperature=temperature,
        max_tokens=max_tokens,
        top_p=1.0,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )
    return resp.choices[0].message.content


def extract_json_from_text(text: str) -> dict:
    """
    Robust JSON extractor:
    - removes ```json fences if present
    - extracts the first {...} block if model adds extra text
    """
    t = text.strip()

    # Remove fenced code blocks
    t = re.sub(r"^```(?:json)?\s*", "", t)
    t = re.sub(r"\s*```$", "", t)

    if not (t.startswith("{") and t.endswith("}")):
        m = re.search(r"\{.*\}", t, flags=re.DOTALL)
        if not m:
            raise ValueError("Model output did not contain a JSON object.")
        t = m.group(0)

    return json.loads(t)


def save_trade_note_docx(data: dict, ticker: str, asof_date, save_path: str) -> str:
    """
    Writes the trade note dict into a Word document.
    asof_date can be a date or string.
    """
    doc = Document()
    doc.add_heading(f"{ticker} Trade Note ({asof_date})", level=0)

    # Abstract
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(data["Abstract"]["Summary"])

    # Recommendation
    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph("Signal: " + data["Recommendation"]["Signal"])
    doc.add_paragraph("Rationale: " + data["Recommendation"]["Rationale"])

    # Strategy
    doc.add_heading("Strategy", level=1)
    doc.add_paragraph(data["Strategy"]["Overview"])
    doc.add_paragraph("Indicators used: " + data["Strategy"]["Indicators_used"])

    # Backtest Summary
    doc.add_heading("Backtest Summary", level=1)
    doc.add_paragraph(data["Backtest_summary"]["Key_metrics"])
    doc.add_paragraph(data["Backtest_summary"]["Performance_interpretation"])

    # Risks
    doc.add_heading("Limitations and Risks", level=1)
    doc.add_paragraph(data["Limitations_and_risks"]["Discussion"])

    # Future Refinements
    doc.add_heading("Future Refinements", level=1)
    doc.add_paragraph(data["Future_refinements"]["Discussion"])

    doc.save(save_path)
    return save_path


# -----------------------
# Main
# -----------------------
def main(
    client=None,
    ticker: str = "KO",

    eval_start: str = "2016-01-20",
    eval_end: str   = "2026-01-25",
    buffer_start: str = "2015-01-01",

    start_cash: float = START_CASH,
    commission: float = COMMISSION,

    save_csv: bool = True,
    save_json: bool = True,
    run_btpy: bool = True,
):

    # 1) Data 
    df = download_ohlcv(
        ticker=ticker,
        buffer_start=buffer_start,  # e.g. "2015-01-01"
        eval_start=eval_start,      # e.g. "2016-01-20"
        eval_end=eval_end,          # e.g. "2026-01-25"
    )

    if df is None or df.empty:
        raise RuntimeError("No price data downloaded. Cannot run backtest/metrics.")

    # 2) Indicators
    df = compute_indicators(df)

    # 3) Clean rows for signal/backtest
    df = df.dropna(subset=REQUIRED_COLS).copy()

    # 4) Signals/position/recommendation
    df = compute_signals_and_position(df)

    # 5) Strategy equity (Adj Close, net)
    df_strat, strat_metrics = compute_strategy_equity_adj_net(
    df=df,
    start_cash=start_cash,
    commission=commission,
    required_cols=REQUIRED_COLS,
    eval_start=eval_start,
    eval_end=eval_end,
    )       


    # 6) Buy & Hold (Adj Close, net)
    bh_equity, bh_metrics = compute_buy_hold_adj_net(
        df=df,
        start_cash=start_cash,
        commission=commission,
        required_cols=REQUIRED_COLS,
        eval_start=eval_start,
        eval_end=eval_end,
    )
    # 7) Plot graphs
    plt.figure(figsize=(20, 8))
    plt.plot(df_strat.index, df_strat["strategy_adj_equity_net"], label="Strategy (Adj Close, net)")
    plt.plot(bh_equity.index, bh_equity.values, label="Buy & Hold (Adj Close, net)")
    plt.title("Equity Curve (Net of Commission)")
    plt.xlabel("Date")
    plt.ylabel("Equity ($)")
    plt.legend()
    plt.grid(True)
    plt.show()

    # Drawdowns
    bh_roll_max = bh_equity.cummax()
    bh_drawdown = (bh_equity / bh_roll_max) - 1

    strat_equity =  df_strat["strategy_adj_equity_net"]
    strat_roll_max = strat_equity.cummax()
    strat_drawdown = (strat_equity / strat_roll_max) - 1

    plt.figure(figsize=(20, 8))
    plt.plot(bh_drawdown.index, bh_drawdown.values, label="Buy & Hold Drawdown")
    plt.plot(strat_drawdown.index, strat_drawdown.values, label="Strategy Drawdown")
    plt.title("Drawdown (Net of Commission)")
    plt.xlabel("Date")
    plt.ylabel("Drawdown")
    plt.legend()
    plt.grid(True)
    plt.show()

    # BUY/ SELL execution point
    df_strat["exec_change"] = df_strat["position_shifted"].diff().fillna(0)

    buy_exec  = df_strat["exec_change"] == 1
    sell_exec = df_strat["exec_change"] == -1
    plt.figure(figsize=(20, 8))

    # Price line
    plt.plot(
        df_strat.index,
        df_strat["Close"],
        label="Close Price",
        linewidth=2
    )

    # BUY executions
    plt.scatter(
        df_strat.index[buy_exec],
        df_strat.loc[buy_exec, "Close"],
        marker="^",
        color="green",
        s=120,
        label="BUY executed"
    )

    # SELL executions
    plt.scatter(
        df_strat.index[sell_exec],
        df_strat.loc[sell_exec, "Close"],
        marker="v",
        color="red",
        s=120,
        label="SELL executed"
    )

    plt.title("Trade Execution Chart (Next-Day Execution)")
    plt.xlabel("Date")
    plt.ylabel("Price")
    plt.legend()
    plt.grid(True)
    plt.tight_layout()
    plt.show()
    
    # 8) backtesting.py run (optional) for trades table consistency
    trades = pd.DataFrame()
    if run_btpy:
        df_bt = df.copy()
        stats, trades = run_backtesting_py(df_bt, cash=start_cash, commission=commission)

    # 9) Hit rate + n_trades
    # Use evaluation-window dataframe returned by compute_strategy_equity_adj_net
    pos = df_strat["position_shifted"].fillna(0).astype(float)
    px  = df_strat["Adj Close"].astype(float)

    # Entry: 0 -> 1, Exit: 1 -> 0
    entry_mask = (pos.shift(1).fillna(0) == 0) & (pos == 1)
    exit_mask  = (pos.shift(1).fillna(0) == 1) & (pos == 0)

    entry_px = px[entry_mask].values
    exit_px  = px[exit_mask].values

    # Pair trades safely
    n = min(len(entry_px), len(exit_px))

    if n == 0:
        hit_rate_pct = float("nan")
        n_trades = 0
    else:
        entry_px = entry_px[:n]
        exit_px  = exit_px[:n]

        trade_returns_gross = (exit_px / entry_px) - 1
        trade_returns_net   = trade_returns_gross - 2 * commission   # entry + exit

        hit_rate_pct = float((trade_returns_net > 0).mean() * 100)
        n_trades = n

    # 10) Build prompt
    user_prompt = build_user_prompt(df, ticker, strat_metrics, bh_metrics, hit_rate_pct, n_trades)

    # 11) Save outputs locally
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    if save_csv:
        df.to_csv(f"{ticker}_signals_{ts}.csv")
    if save_json:
        with open(f"{ticker}_prompt_inputs_{ts}.json", "w") as f:
            json.dump(
                {
                    "ticker": ticker,
                    "date": str(df.index[-1].date()),
                    "strat_metrics": strat_metrics,
                    "bh_metrics": bh_metrics,
                    "hit_rate_pct": hit_rate_pct,
                    "n_trades": n_trades,
                    "prompt": user_prompt,
                },
                f,
                indent=2,
            )
    system_prompt = (
        "You are a professional technical analyst."
        "You are given pre-computed indicators, signals, and backtest metrics."
        "Do NOT invent numbers or assumptions."
        "Base your analysis strictly on the provided inputs."
    )
    # 12) LLM trade note 
    ai_text = generate_trade_note_llm(
        client=client,
        deployment=os.environ["AZURE_DEPLOYMENT_NAME"],
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        temperature=0.2,
        max_tokens=800,
    )

    print("=== Trade Note (raw) ===")
    print(ai_text)

    trade_note = extract_json_from_text(ai_text)

    # 13) Save DOCX 
    asof_date = df.index[-1].date()
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    save_path = f"{ticker}_trade_note_{asof_date}_{ts}.docx"   

    doc_path = save_trade_note_docx(trade_note, ticker=ticker, asof_date=asof_date, save_path=save_path)
    print("Saved to:", doc_path)
    print(bh_equity)
    print(strat_metrics)
    print(bh_metrics)

    return df, df_strat, bh_equity, strat_metrics, bh_metrics, user_prompt, trades, trade_note, doc_path


if __name__ == "__main__":
    client = build_azure_client()

    (
        df,
        df_strat,
        bh_equity,
        strat_metrics,
        bh_metrics,
        user_prompt,
        trades,
        trade_note,
        doc_path,
    ) = main(client=client)
    

    print("Trade note saved to:", doc_path)
