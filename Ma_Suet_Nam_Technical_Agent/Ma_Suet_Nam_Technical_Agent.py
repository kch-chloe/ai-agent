#!/usr/bin/env python
# coding: utf-8

# In[408]:


# Cell 0 — Imports & Desktop output folder

import os
import json
from pathlib import Path
from datetime import datetime

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import yfinance as yf
import backtrader as bt

from dotenv import load_dotenv
load_dotenv()

# Optional: docx export
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_OK = True
except Exception:
    DOCX_OK = False

RUN_TS = datetime.now().strftime("%Y%m%d_%H%M%S")
ROOT = Path("..") / "OUTPUT" / f"KO_TECH_AGENT_{RUN_TS}"

DIRS = {
    "root": ROOT,
    "data": ROOT / "data",
    "outputs": ROOT / "outputs",
    "plots": ROOT / "plots",
    "llm": ROOT / "llm",
}

for p in DIRS.values():
    p.mkdir(parents=True, exist_ok=True)

print("✅ Output folder created", ROOT)

# In[409]:


# Cell 1 — Parameters 

TICKER = "KO"
PERIOD = "11y"
INTERVAL = "1d"

# Strategy params
FAST_P = 25
BOLL_TIMEPERIOD = 10
BIAS_YUZHI1 = -1
BIAS_YUZHI2 = 5

STOP_LOSS = 70       # /1000
STOP_PROFIT = 135    # /1000
CW_BILI = 0.75

START_EQUITY = 10000
COMMISSION = 0.001
RISK_FREE_RATE = 0.0

AUTO_ADJUST = False

AZURE_OPENAI_ENDPOINT = os.environ["AZURE_ENDPOINT"]
AZURE_OPENAI_API_KEY = os.environ["AZURE_API_KEY"]
AZURE_OPENAI_API_VERSION = os.environ.get("AZURE_API_VERSION", "2024-02-15-preview")
AZURE_OPENAI_DEPLOYMENT = os.environ["AZURE_DEPLOYMENT_NAME"]

print("✅ Params loaded.")
print("Desktop folder:", ROOT)
print("Azure key present?", bool(AZURE_OPENAI_API_KEY))

# In[410]:


# Cell 2 — Download + clean OHLCV

raw = yf.download(
    TICKER,
    period=PERIOD,
    interval=INTERVAL,
    auto_adjust=AUTO_ADJUST,
    progress=False
)

df = raw.copy()

# Flatten MultiIndex columns if needed
if isinstance(df.columns, pd.MultiIndex):
    df.columns = [c[0] if c[1] == "" else c[0] for c in df.columns]

df = df.reset_index()

dt_col = None
for c in ["Date", "Datetime", "date", "datetime"]:
    if c in df.columns:
        dt_col = c
        break
if dt_col is None:
    raise ValueError(f"Cannot find datetime column. Columns: {df.columns.tolist()}")

df.rename(columns={dt_col: "datetime"}, inplace=True)

rename_map = {
    "Open": "open",
    "High": "high",
    "Low": "low",
    "Close": "close",
    "Adj Close": "adj close",
    "Volume": "volume",
}
df.rename(columns=rename_map, inplace=True)

needed = ["datetime", "open", "high", "low", "close", "adj close", "volume"]
for col in needed:
    if col not in df.columns:
        if col == "adj close" and "close" in df.columns:
            df["adj close"] = df["close"]
        else:
            raise ValueError(f"Missing column '{col}'. Available: {df.columns.tolist()}")

df = df[needed].copy()
df["datetime"] = pd.to_datetime(df["datetime"], errors="coerce")
df = df.dropna(subset=["datetime"]).copy()

for c in ["open","high","low","close","adj close","volume"]:
    df[c] = pd.to_numeric(df[c], errors="coerce")

df = df.dropna(subset=["open","high","low","close","adj close"]).copy()
df = df.set_index("datetime").sort_index()

p_raw = DIRS["data"] / f"{TICKER}_{PERIOD}_{INTERVAL}_raw.csv"
p_clean = DIRS["data"] / f"{TICKER}_{PERIOD}_{INTERVAL}_clean_auto_adjust_{AUTO_ADJUST}.csv"
raw.to_csv(p_raw)
df.to_csv(p_clean)

print("✅ Saved raw:", p_raw)
print("✅ Saved clean:", p_clean)
print("df shape:", df.shape)
df.tail()

# In[411]:


# Cell 3 — Indicators (SMA + Bollinger + Bias)

def calc_bollinger(close_list, timeperiod=10, nbdevup=2, nbdevdn=2):
    close = np.array(close_list, dtype=float)
    mid = np.full_like(close, np.nan, dtype=float)
    up  = np.full_like(close, np.nan, dtype=float)
    dn  = np.full_like(close, np.nan, dtype=float)

    for i in range(timeperiod - 1, len(close)):
        w = close[i - timeperiod + 1 : i + 1]
        ma = np.mean(w)
        std = np.std(w, ddof=1)  # sample std
        mid[i] = ma
        up[i]  = ma + nbdevup * std
        dn[i]  = ma - nbdevdn * std
    return up, mid, dn

PRICE_COL_SIGNAL = "close"  # keep consistent with your backtest if needed
PRICE_COL_EXEC = "close"

df["fast_ma"] = df[PRICE_COL_SIGNAL].rolling(FAST_P).mean()

boll_up, boll_mid, boll_dn = calc_bollinger(
    close_list=df[PRICE_COL_SIGNAL].values,
    timeperiod=BOLL_TIMEPERIOD,
    nbdevup=2,
    nbdevdn=2
)
df["boll_up"] = boll_up
df["boll_mid"] = boll_mid
df["boll_dn"] = boll_dn

df["bias"] = (df[PRICE_COL_SIGNAL] - df["fast_ma"]) / df["fast_ma"] * 100.0

# Save features table
features_cols = ["open","high","low","close","adj close","volume","fast_ma","boll_up","boll_mid","boll_dn","bias"]
p_feat = DIRS["data"] / "KO_features.csv"
df[features_cols].to_csv(p_feat)

print("✅ Saved features:", p_feat)
df.tail()

# In[412]:


# Cell 4 — Strategy signal 

signal_price = df[PRICE_COL_SIGNAL]
pre_signal_price = df[PRICE_COL_SIGNAL].shift(1)

df["pre_boll_mid"] = df["boll_mid"].shift(1)

df["cross_up_mid"] = (signal_price > df["boll_mid"]) & (pre_signal_price <= df["pre_boll_mid"])
df["cond_bias_low"] = df["bias"] <= BIAS_YUZHI1
df["cond_bias_high"] = df["bias"] >= BIAS_YUZHI2
df["cond_below_ma"] = signal_price < df["fast_ma"]

# IMPORTANT: Replace this EXACTLY with your real entry/exit logic if your backtest uses different rules.
df["entry_signal"] = df["cross_up_mid"] & df["cond_bias_low"] & df["cond_below_ma"]
df["exit_signal"] = df["cond_bias_high"]

def decide(row):
    if bool(row["exit_signal"]):
        return "SELL"
    if bool(row["entry_signal"]):
        return "BUY"
    return "HOLD"

df["recommendation"] = df.apply(decide, axis=1)

# Export tables asked in your workflow
p_full = DIRS["outputs"] / "indicator_table_full.csv"
p_60 = DIRS["outputs"] / "indicator_table_last_60d.csv"
df.to_csv(p_full)
df.tail(60).to_csv(p_60)

print("✅ Saved:", p_full)
print("✅ Saved:", p_60)

df[["close","adj close","fast_ma","boll_mid","bias","entry_signal","exit_signal","recommendation"]].tail(10)

# In[413]:


# Cell 5 — Backtest + metrics JSON  (+ exit time & reason export)

money = []

class KOExactStrategy(bt.Strategy):
    params = (
        ("stop_loss", STOP_LOSS),
        ("stop_profit", STOP_PROFIT),
    )

    def __init__(self):
        self.sl = self.p.stop_loss / 1000.0
        self.tp = self.p.stop_profit / 1000.0

        self.entry_order = None
        self.exit_order = None
        self.stop_loss_order = None
        self.stop_profit_order = None
        self.stop_loss_price = None
        self.stop_profit_price = None

        self.closed_trade_pnls_comm = []

        # no-leakage audit
        self._pending_signal_dt = None
        self.bad_samebar_fills = []

        # NEW: pending exit reason for market close
        self._pending_exit_reason = None

        # NEW: exit events list
        self.exit_events = []

    def _cancel_exit_orders(self):
        if self.stop_loss_order is not None:
            self.cancel(self.stop_loss_order)
            self.stop_loss_order = None
        if self.stop_profit_order is not None:
            self.cancel(self.stop_profit_order)
            self.stop_profit_order = None

    def notify_trade(self, trade):
        if trade.isclosed:
            self.closed_trade_pnls_comm.append(float(trade.pnlcomm))

    def notify_order(self, order):
        if order.status in [order.Submitted, order.Accepted]:
            return

        if order.status == order.Completed:
            exec_dt = bt.num2date(order.executed.dt).replace(tzinfo=None)

            if order.isbuy():
                exec_price = float(order.executed.price)
                exec_size = float(order.executed.size)

                # same-bar fill audit
                if self._pending_signal_dt is not None and exec_dt <= self._pending_signal_dt:
                    self.bad_samebar_fills.append((self._pending_signal_dt, exec_dt, exec_price))
                self._pending_signal_dt = None

                self.stop_loss_price = exec_price * (1 - self.sl)
                self.stop_profit_price = exec_price * (1 + self.tp)

                self._cancel_exit_orders()

                # OCO: Stop + Limit
                self.stop_loss_order = self.sell(
                    exectype=bt.Order.Stop,
                    price=self.stop_loss_price,
                    size=exec_size,
                    transmit=False
                )
                self.stop_profit_order = self.sell(
                    exectype=bt.Order.Limit,
                    price=self.stop_profit_price,
                    size=exec_size,
                    oco=self.stop_loss_order,
                    transmit=True
                )

                self.entry_order = None

            elif order.issell():
                # determine exit reason
                reason = "OTHER"
                if order == self.stop_loss_order:
                    reason = "SL"
                elif order == self.stop_profit_order:
                    reason = "TP"
                elif order == self.exit_order:
                    reason = self._pending_exit_reason or "SIGNAL"

                self.exit_events.append({
                    "exit_dt": exec_dt,
                    "reason": reason,
                })

                # reset
                self._pending_exit_reason = None
                self.entry_order = None
                self.exit_order = None
                self.stop_loss_order = None
                self.stop_profit_order = None
                self.stop_loss_price = None
                self.stop_profit_price = None

        elif order.status in [order.Canceled, order.Margin, order.Rejected]:
            if order == self.entry_order:
                self.entry_order = None
                self._pending_signal_dt = None
            if order == self.exit_order:
                self.exit_order = None
                self._pending_exit_reason = None
            if order == self.stop_loss_order:
                self.stop_loss_order = None
            if order == self.stop_profit_order:
                self.stop_profit_order = None

    def next(self):
        i = len(self.data) - 1
        if i < FAST_P or i >= len(df):
            return

        if self.entry_order is not None or self.exit_order is not None:
            return

        money.append(self.broker.get_value())

        signal_price = float(df[PRICE_COL_SIGNAL].iloc[i])
        pre_signal_price = float(df[PRICE_COL_SIGNAL].iloc[i - 1])

        sizing_price = float(df["close"].iloc[i])

        boll_mid = float(df["boll_mid"].iloc[i])
        pre_boll_mid = float(df["boll_mid"].iloc[i - 1])
        bias = float(df["bias"].iloc[i])
        fast_ma = float(df["fast_ma"].iloc[i])

        cross_up_mid = (signal_price > boll_mid) and (pre_signal_price <= pre_boll_mid)
        cond_bias_low = bias <= BIAS_YUZHI1
        cond_bias_high = bias >= BIAS_YUZHI2
        cond_below_ma = signal_price < fast_ma

        size = int((self.broker.get_value() / sizing_price) * CW_BILI)

        if self.position:
            # hard SL fuse
            if self.stop_loss_price is not None:
                cur_low = float(self.data.low[0])
                if cur_low <= float(self.stop_loss_price):
                    self._cancel_exit_orders()
                    self._pending_exit_reason = "SL_FUSE"
                    self.exit_order = self.close()
                    return

            if cond_bias_high:
                self._cancel_exit_orders()
                self._pending_exit_reason = "SIGNAL"
                self.exit_order = self.close()
            return

        if cross_up_mid and cond_bias_low and cond_below_ma and size > 0:
            self._pending_signal_dt = self.data.datetime.datetime(0)
            self.entry_order = self.buy(size=size)


# Run backtest
money.clear()

cerebro = bt.Cerebro()
cerebro.broker.set_coc(False)

data = bt.feeds.PandasData(
    dataname=df,
    datetime=None,
    open='open',
    high='high',
    low='low',
    close='close',
    volume='volume',
    openinterest=-1,
    timeframe=bt.TimeFrame.Days,
)
cerebro.adddata(data)
cerebro.addstrategy(KOExactStrategy)

cerebro.broker.setcash(START_EQUITY)
cerebro.broker.setcommission(commission=COMMISSION)

st = cerebro.run()[0]

# audit print
if len(st.bad_samebar_fills) > 0:
    print("⚠️ WARNING: Found same-bar fills (signal_dt, exec_dt, exec_price):")
    for x in st.bad_samebar_fills[:10]:
        print(x)
else:
    print("✅ No same-bar fills detected (orders filled after signal bar).")

final_value = float(cerebro.broker.getvalue())
total_ret_pct = (final_value / START_EQUITY - 1) * 100.0

# equity curve aligned from FAST_P
start_i = FAST_P
end_i = min(len(df), start_i + len(money))
eq_idx = df.index[start_i:end_i]
n = min(len(eq_idx), len(money))
eq_idx = eq_idx[:n]
strategy_equity = pd.Series(money[:n], index=eq_idx, name="Strategy")

def annualized_sharpe(eq: pd.Series, rf: float = 0.0, periods: int = 252) -> float:
    r = eq.pct_change().dropna()
    if len(r) == 0 or r.std() == 0:
        return float("nan")
    return float(((r.mean() * periods) - rf) / (r.std() * np.sqrt(periods)))

def max_drawdown_pct(eq: pd.Series) -> float:
    peak = eq.cummax()
    dd = (eq / peak - 1.0) * 100.0
    return float(dd.min())

def cagr(eq: pd.Series) -> float:
    eq = eq.dropna()
    if len(eq) < 2:
        return float("nan")
    days = (eq.index[-1] - eq.index[0]).days
    years = days / 365.25
    if years <= 0:
        return float("nan")
    return float((eq.iloc[-1] / eq.iloc[0]) ** (1.0 / years) - 1.0)

# hit rate
pnls = np.array(st.closed_trade_pnls_comm, dtype=float)
total_closed = int(len(pnls))
wins = int((pnls > 0).sum())
hit_rate_pct = (wins / total_closed * 100.0) if total_closed > 0 else None

print("✅ Closed trades:", total_closed, "| Wins:", wins, "| Hit rate %:", hit_rate_pct)

# export EXIT_EVENTS.csv
exit_df = pd.DataFrame(st.exit_events)
if len(exit_df) > 0:
    exit_df["exit_dt"] = pd.to_datetime(exit_df["exit_dt"])
    exit_df.sort_values("exit_dt", inplace=True)

p_exit = DIRS["outputs"] / "EXIT_EVENTS.csv"
exit_df.to_csv(p_exit, index=False, encoding="utf-8")
print("✅ Saved exit events:", p_exit)

# metrics JSON
metrics = {
    "ticker": TICKER,
    "asof": str(df.index[-1].date()),
    "params": {
        "FAST_P": FAST_P,
        "BOLL_TIMEPERIOD": BOLL_TIMEPERIOD,
        "BIAS_YUZHI1": BIAS_YUZHI1,
        "BIAS_YUZHI2": BIAS_YUZHI2,
        "STOP_LOSS": STOP_LOSS,
        "STOP_PROFIT": STOP_PROFIT,
        "CW_BILI": CW_BILI,
        "START_EQUITY": START_EQUITY,
        "COMMISSION": COMMISSION,
        "AUTO_ADJUST": AUTO_ADJUST,
        "PRICE_COL_SIGNAL": PRICE_COL_SIGNAL,
        "PRICE_COL_EXEC": PRICE_COL_EXEC,
    },
    "strategy": {
        "final_value": final_value,
        "total_return_pct": float(total_ret_pct),
        "cagr": float(cagr(strategy_equity)),
        "sharpe": float(annualized_sharpe(strategy_equity, rf=RISK_FREE_RATE)),
        "max_drawdown_pct": float(max_drawdown_pct(strategy_equity)),
        "hit_rate_pct": hit_rate_pct,
        "total_trades_closed": total_closed,
        "won_trades": wins,
    }
}

p_metrics = DIRS["outputs"] / "BACKTEST_METRICS.json"
with open(p_metrics, "w", encoding="utf-8") as f:
    json.dump(metrics, f, ensure_ascii=False, indent=2)

print("✅ Saved metrics:", p_metrics)
print("Total return %:", metrics["strategy"]["total_return_pct"])

# In[414]:


# Cell 5.5— Plot EXECUTED buy/sell points (Backtrader) + print first trade return

import matplotlib.pyplot as plt
import pandas as pd
import backtrader as bt

class StrategyWithExec(KOExactStrategy):
    def __init__(self):
        super().__init__()
        self.exec_events = []  # list of dicts

    def notify_order(self, order):
        super().notify_order(order)

        if order.status == order.Completed:
            dt = self.data.datetime.datetime(0)
            px = float(order.executed.price)
            side = "BUY" if order.isbuy() else "SELL" if order.issell() else "OTHER"

            reason = None
            if side == "SELL":
                if getattr(self, "stop_profit_order", None) is not None and order.ref == self.stop_profit_order.ref:
                    reason = "TP"
                elif getattr(self, "stop_loss_order", None) is not None and order.ref == self.stop_loss_order.ref:
                    reason = "SL"
                else:
                    reason = "SIGNAL"

            self.exec_events.append({
                "dt": dt,
                "side": side,
                "price": px,
                "reason": reason
            })

# rerun a lightweight backtest to capture exec points
c = bt.Cerebro()
d = bt.feeds.PandasData(
    dataname=df,
    datetime=None,
    open='open', high='high', low='low', close='close',
    volume='volume',
    openinterest=-1,
    timeframe=bt.TimeFrame.Days,
)
c.adddata(d)
c.addstrategy(StrategyWithExec)
c.broker.setcash(START_EQUITY)
c.broker.setcommission(commission=COMMISSION)

st_exec = c.run()[0]
ev = pd.DataFrame(st_exec.exec_events)

# split buys/sells
buys = ev[ev["side"] == "BUY"].copy()
sells = ev[ev["side"] == "SELL"].copy()

# plot on execution price series (close)
price_series = df[PRICE_COL_EXEC].astype(float)

plt.figure(figsize=(14, 6))
plt.plot(df.index, price_series.values, label=f"Price ({PRICE_COL_EXEC})")

if len(buys) > 0:
    plt.scatter(buys["dt"], buys["price"], marker="^", s=80, label="BUY (executed)")
if len(sells) > 0:
    plt.scatter(sells["dt"], sells["price"], marker="v", s=80, label="SELL (executed)")

plt.title(f"{TICKER} Executed Buy/Sell Points (Backtrader)")
plt.xlabel("Date"); plt.ylabel("Price")
plt.grid(True); plt.legend()

p_exec = DIRS["plots"] / "buy_sell_points_executed.png"
plt.savefig(p_exec, bbox_inches="tight")
# plt.show()

print("✅ Saved:", p_exec)
print("Executed buys:", len(buys), "| Executed sells:", len(sells))

# compute first trade return (pair 1st buy with 1st sell after it)
if len(buys) > 0 and len(sells) > 0:
    b0 = buys.iloc[0]
    s0 = sells[sells["dt"] > b0["dt"]].iloc[0]
    ret = (s0["price"] / b0["price"] - 1) * 100
    print("\nFirst trade:")
    print(" Entry:", b0["dt"], "price:", b0["price"])
    print(" Exit :", s0["dt"], "price:", s0["price"], "reason:", s0.get("reason"))
    print(f" Return: {ret:.2f}% (TP target ~ {STOP_PROFIT/10:.1f}%)")

# In[415]:


# Cell 6 — Plots (equity + drawdown)  

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

def annualized_sharpe(eq: pd.Series, rf: float = 0.0, periods: int = 252) -> float:
    r = eq.pct_change().dropna()
    if len(r) == 0 or r.std() == 0:
        return float("nan")
    return float(((r.mean() * periods) - rf) / (r.std() * np.sqrt(periods)))

def drawdown_series(eq: pd.Series) -> pd.Series:
    peak = eq.cummax()
    return (eq / peak - 1.0) * 100.0

# ----------------------------
# Buy & Hold equity curve (Adj Close) aligned to strategy period
# ----------------------------
if "adj close" not in df.columns:
    raise ValueError("No 'adj close' in df，cant cal Buy & Hold。")

bh_prices = df["adj close"].reindex(strategy_equity.index).astype(float).ffill()

if bh_prices.isna().any():
    raise ValueError("there is N/a in adj close")

first_price = float(bh_prices.iloc[0])
if first_price <= 0:
    raise ValueError("adj<0")

# can have less than one share
bh_shares = START_EQUITY / first_price
buyhold_equity = pd.Series(bh_shares * bh_prices, index=bh_prices.index, name="Buy & Hold (Adj Close)")

# Metrics for legend
sh_strategy = annualized_sharpe(strategy_equity)
sh_bh = annualized_sharpe(buyhold_equity)

dd_strategy = drawdown_series(strategy_equity)
dd_bh = drawdown_series(buyhold_equity)

# Plot: Equity / Return curve (Strategy vs Buy&Hold) + Sharpe in legend
plt.figure(figsize=(12, 5))
plt.plot(strategy_equity.index, strategy_equity.values,
         label=f"Strategy Equity (Sharpe={sh_strategy:.3f})")
plt.plot(buyhold_equity.index, buyhold_equity.values,
         label=f"Buy & Hold (Adj Close) (Sharpe={sh_bh:.3f})")
plt.title("Equity Curve (Strategy vs Buy & Hold)")
plt.xlabel("Date")
plt.ylabel("Portfolio Value")
plt.grid(True)
plt.legend()

p_eq = DIRS["plots"] / "equity_curve.png"
plt.savefig(p_eq, bbox_inches="tight")
# plt.show()

# Plot: Drawdown (both)
plt.figure(figsize=(12, 5))
plt.plot(dd_strategy.index, dd_strategy.values, label="Strategy Drawdown (%)")
plt.plot(dd_bh.index, dd_bh.values, label="Buy & Hold (Adj Close) Drawdown (%)")
plt.title("Drawdown (Strategy vs Buy & Hold)")
plt.xlabel("Date")
plt.ylabel("Drawdown (%)")
plt.grid(True)
plt.legend()

p_dd = DIRS["plots"] / "drawdown.png"
plt.savefig(p_dd, bbox_inches="tight")
# plt.show()

print("✅ Saved:", p_eq)
print("✅ Saved:", p_dd)
print(f"Sharpe — Strategy: {sh_strategy:.4f} | Buy & Hold (Adj Close): {sh_bh:.4f}")

# In[416]:


# Cell 7 — Azure OpenAI Trade Note (minimal layout fix: NO bullets; prompt + renderer only)

import json
import re
import pandas as pd
import numpy as np
from pathlib import Path
from openai import AzureOpenAI

# 0) Guardrails
if "df" not in globals() or df is None or len(df) < 5:
    raise ValueError("df is missing or too short. Run the data/indicator cells first.")

required_cols = ["close", "adj close", "fast_ma", "boll_mid", "boll_up", "boll_dn", "bias"]
missing_cols = [c for c in required_cols if c not in df.columns]
if missing_cols:
    raise ValueError(f"df missing required columns: {missing_cols}")

if "TICKER" not in globals():
    raise NameError("Missing TICKER. Run your parameter cell first.")
if "FAST_P" not in globals():
    raise NameError("Missing FAST_P. Run your parameter cell first.")

if not AZURE_OPENAI_API_KEY:
    raise ValueError("Missing AZURE_OPENAI_API_KEY in environment variables.")
if not AZURE_OPENAI_ENDPOINT or not AZURE_OPENAI_API_VERSION or not AZURE_OPENAI_DEPLOYMENT:
    raise ValueError("Missing Azure OpenAI config variables (endpoint/version/deployment).")

client = AzureOpenAI(
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
)

# 1) Locate outputs folder (align with your pipeline)
if "DIRS" in globals() and isinstance(DIRS, dict) and "outputs" in DIRS and "llm" in DIRS:
    out_dir = Path(DIRS["outputs"])
    llm_dir = Path(DIRS["llm"])
else:
    out_dir = Path("./outputs")
    llm_dir = Path("./llm")

out_dir.mkdir(parents=True, exist_ok=True)
llm_dir.mkdir(parents=True, exist_ok=True)

p_metrics = out_dir / "BACKTEST_METRICS.json"
p_exit = out_dir / "EXIT_EVENTS.csv"

# 2) Helper functions
def fmt(x, nd=2):
    try:
        return f"{float(x):.{nd}f}"
    except Exception:
        return "Not available"

def fmt_int(x):
    try:
        return str(int(x))
    except Exception:
        return "Not available"

def annualized_sharpe(eq: pd.Series, rf: float = 0.0, periods: int = 252) -> float:
    r = eq.pct_change().dropna()
    if len(r) == 0 or r.std() == 0:
        return float("nan")
    return float(((r.mean() * periods) - rf) / (r.std() * np.sqrt(periods)))

def max_drawdown_pct(eq: pd.Series) -> float:
    peak = eq.cummax()
    dd = (eq / peak - 1.0) * 100.0
    return float(dd.min())

def cagr(eq: pd.Series) -> float:
    eq = eq.dropna()
    if len(eq) < 2:
        return float("nan")
    years = (eq.index[-1] - eq.index[0]).days / 365.25
    if years <= 0:
        return float("nan")
    return float((eq.iloc[-1] / eq.iloc[0]) ** (1.0 / years) - 1.0)

def total_return_pct(eq: pd.Series) -> float:
    eq = eq.dropna()
    if len(eq) < 2:
        return float("nan")
    return float((eq.iloc[-1] / eq.iloc[0] - 1.0) * 100.0)

def build_bh_equity(prices: pd.Series, start_equity: float) -> pd.Series:
    prices = prices.astype(float).ffill()
    shares = start_equity / float(prices.iloc[0])
    return shares * prices

# 3) Load strategy metrics (prefer existing globals; else fallback to BACKTEST_METRICS.json)
m = {}
if p_metrics.exists():
    with open(p_metrics, "r", encoding="utf-8") as f:
        m = json.load(f)
st = m.get("strategy", {}) if isinstance(m, dict) else {}

strategy_total_return_pct = globals().get("strategy_total_return_pct", st.get("total_return_pct", None))
strategy_cagr_pct        = globals().get("strategy_cagr_pct",        st.get("cagr_pct", None))
strategy_sharpe          = globals().get("strategy_sharpe",          st.get("sharpe", None))
strategy_max_drawdown_pct= globals().get("strategy_max_drawdown_pct",st.get("max_drawdown_pct", None))

# reconstruct strategy CAGR if missing
if strategy_cagr_pct is None and strategy_total_return_pct is not None:
    try:
        bt_start = pd.to_datetime(st.get("bt_start"))
        bt_end   = pd.to_datetime(st.get("bt_end"))
        years = (bt_end - bt_start).days / 365.25
        if years > 0:
            total_mult = 1.0 + float(strategy_total_return_pct) / 100.0
            strategy_cagr_pct = (total_mult ** (1.0 / years) - 1.0) * 100.0
    except Exception:
        pass

hit_rate = globals().get("hit_rate", st.get("hit_rate_pct", None))

# trades
trade_returns = globals().get("trade_returns", None)
num_trades = None
if trade_returns is not None:
    try:
        num_trades = len(trade_returns)
    except Exception:
        num_trades = None
if num_trades is None:
    num_trades = st.get("total_trades_closed", None)
if num_trades is None and p_exit.exists():
    try:
        ex = pd.read_csv(p_exit)
        num_trades = len(ex)
    except Exception:
        pass

# 4) Compute B&H metrics if missing (from df)
bh_total_return_pct = globals().get("bh_total_return_pct", None)
bh_cagr_pct         = globals().get("bh_cagr_pct", None)
bh_sharpe           = globals().get("bh_sharpe", None)
bh_max_drawdown_pct = globals().get("bh_max_drawdown_pct", None)

eval_df = df.iloc[int(FAST_P):].copy()
start_equity = float(m.get("start_equity", globals().get("START_EQUITY", 10000))) if isinstance(m, dict) else float(globals().get("START_EQUITY", 10000))
rf = float(m.get("risk_free_rate", globals().get("RISK_FREE_RATE", 0.0))) if isinstance(m, dict) else float(globals().get("RISK_FREE_RATE", 0.0))

if any(v is None for v in [bh_total_return_pct, bh_cagr_pct, bh_sharpe, bh_max_drawdown_pct]):
    eq_bh = build_bh_equity(eval_df["adj close"], start_equity=start_equity)
    bh_total_return_pct = total_return_pct(eq_bh)
    bh_cagr_pct = cagr(eq_bh) * 100.0
    bh_sharpe = annualized_sharpe(eq_bh, rf=rf)
    bh_max_drawdown_pct = max_drawdown_pct(eq_bh)

# diagnostic B&H close
eq_bh_close = build_bh_equity(eval_df["close"], start_equity=start_equity)
bh_close_total_return_pct = total_return_pct(eq_bh_close)
bh_close_cagr_pct = cagr(eq_bh_close) * 100.0
bh_close_sharpe = annualized_sharpe(eq_bh_close, rf=rf)
bh_close_max_drawdown_pct = max_drawdown_pct(eq_bh_close)

# 5) Latest snapshot
ticker = TICKER
asof_date = df.index[-1].date()
last = df.iloc[-1]

# 6) PROMPT 
ser_prompt = user_prompt = f"""
Context: Analysis of the trading strategy for {ticker}.
Write a structured, professional trade note in clean paragraphs. Do NOT use bullet points anywhere.

Ticker: {ticker}
Date: {asof_date}

Non-negotiable disclosure (must appear verbatim):
Strategy backtest/execution uses CLOSE (unadjusted) to avoid adjustment-related leakage in trade simulation.
Buy & Hold benchmark uses ADJ CLOSE as a dividend-inclusive total-return proxy.

Current recommendation (deterministic engine): {str(last.get("recommendation","Not available")).upper()}
Entry signal today: {bool(last.get("entry_signal", False))}
Exit signal today:  {bool(last.get("exit_signal", False))}

Latest indicators:
fast_ma={fmt(last.get("fast_ma"), 4)}, boll_mid={fmt(last.get("boll_mid"), 4)}, boll_up={fmt(last.get("boll_up"), 4)}, boll_dn={fmt(last.get("boll_dn"), 4)}, bias={fmt(last.get("bias"), 4)}.
Close={fmt(last.get("close"), 2)} and Adj Close={fmt(last.get("adj close"), 2)}.

Backtest metrics (net of commission):
Strategy (Close): total return {fmt(strategy_total_return_pct,2)}%, CAGR {fmt(strategy_cagr_pct,2)}%, Sharpe {fmt(strategy_sharpe,4)}, max drawdown {fmt(strategy_max_drawdown_pct,2)}%, hit rate {fmt(hit_rate,2)}%, trades {fmt_int(num_trades)}.
Buy & Hold (Adj Close): total return {fmt(bh_total_return_pct,2)}%, CAGR {fmt(bh_cagr_pct,2)}%, Sharpe {fmt(bh_sharpe,4)}, max drawdown {fmt(bh_max_drawdown_pct,2)}%.
Diagnostic Buy & Hold (Close): total return {fmt(bh_close_total_return_pct,2)}%, CAGR {fmt(bh_close_cagr_pct,2)}%, Sharpe {fmt(bh_close_sharpe,4)}, max drawdown {fmt(bh_close_max_drawdown_pct,2)}%.

Strategy rules (do not modify):
Entry triggers when cross_up_mid is true and bias <= BIAS_YUZHI1 and Close < fast_ma.
Exit can occur via take-profit, stop-loss, or signal exit; signal exit occurs when bias >= BIAS_YUZHI2.
TP is entry_price*(1+STOP_PROFIT/1000), SL is entry_price*(1-STOP_LOSS/1000), and SL_FUSE is an emergency exit when intrabar low breaches the stop-loss threshold.

Writing requirements:
Use short paragraphs with academic-but-practical tone.
Explain the market state using fast_ma, Bollinger mid and bias.
Explain the indicator confluence as a pullback/mean-reversion entry rationale.
State the recommendation (BUY/SELL/HOLD) consistent with deterministic engine.
Compare risk-adjusted performance versus Buy & Hold (Adj Close), focusing on Sharpe and drawdown.
Discuss key limitations and propose future refinements.
Do not invent any numbers, do not use external news, and do not forecast.

Output rule:
Return ONLY a valid JSON object, with each field written as paragraphs only (no bullets, no numbering).

Expected JSON structure:
{{
  "Abstract": {{"Summary": "Max 90 words, one paragraph."}},
  "Recommendation": {{"Signal": "Buy | Sell | Hold", "Rationale": "One short paragraph."}},
  "Strategy": {{"Overview": "Two short paragraphs.", "Indicators_used": "One paragraph."}},
  "Backtest_summary": {{"Key_metrics": "One compact paragraph containing the numbers exactly.", "Performance_interpretation": "One paragraph."}},
  "Limitations_and_risks": {{"Discussion": "One to two paragraphs."}},
  "Future_refinements": {{"Discussion": "One paragraph."}}
}}
"""

# 7) Call Azure OpenAI
resp = client.chat.completions.create(
    model=AZURE_OPENAI_DEPLOYMENT,
    messages=[
        {"role": "system", "content": ser_prompt},
        {"role": "user", "content": user_prompt},
    ],
    temperature=0.2,
)
raw = (resp.choices[0].message.content or "").strip()

# 8) Parse JSON + save 
def _strip_fences(t: str) -> str:
    t = t.strip()
    t = re.sub(r"^```(?:json)?\s*", "", t, flags=re.IGNORECASE)
    t = re.sub(r"\s*```$", "", t)
    return t.strip()

def _try_parse_json(t: str):
    t2 = _strip_fences(t)
    try:
        return json.loads(t2), t2
    except Exception:
        m = re.search(r"\{.*\}", t2, flags=re.DOTALL)
        if m:
            try:
                return json.loads(m.group(0)), m.group(0)
            except Exception:
                pass
    return None, t2

obj, cleaned = _try_parse_json(raw)

p_llm_json = llm_dir / "llm_decision.json"
p_llm_md = llm_dir / "llm_decision.md"

with open(p_llm_json, "w", encoding="utf-8") as f:
    if obj is not None:
        json.dump(obj, f, ensure_ascii=False, indent=2)
    else:
        f.write(cleaned)

if obj is None:
    pretty = cleaned
else:
    def _get(d, *keys, default=""):
        cur = d
        for k in keys:
            if not isinstance(cur, dict) or k not in cur:
                return default
            cur = cur[k]
        return cur if cur is not None else default

    pretty = "\n\n".join([
        "Abstract\n" + _get(obj, "Abstract", "Summary", "").strip(),
        "Recommendation\n" + ("Signal: " + _get(obj, "Recommendation", "Signal", "").strip() + ". " +
                              _get(obj, "Recommendation", "Rationale", "").strip()).strip(),
        "Strategy\n" + _get(obj, "Strategy", "Overview", "").strip(),
        "Indicators used\n" + _get(obj, "Strategy", "Indicators_used", "").strip(),
        "Backtest Summary\n" + (_get(obj, "Backtest_summary", "Key_metrics", "").strip() + "\n" +
                                _get(obj, "Backtest_summary", "Performance_interpretation", "").strip()).strip(),
        "Limitations and Risks\n" + _get(obj, "Limitations_and_risks", "Discussion", "").strip(),
        "Future Refinements\n" + _get(obj, "Future_refinements", "Discussion", "").strip(),
    ]).strip()

with open(p_llm_md, "w", encoding="utf-8") as f:
    f.write(pretty)

print("✅ LLM decision saved:")
print(" -", p_llm_json)
print(" -", p_llm_md)
print("\n---- Preview (first 600 chars) ----")

# In[417]:


# Cell 8 Build TRADENOTE — friend format + realised exits + TOTAL exits

import json, re
import pandas as pd
import numpy as np
from pathlib import Path

REPORT_BASENAME = "TRADENOTE"

# ---------- Paths ----------
if "DIRS" in globals() and isinstance(DIRS, dict) and "outputs" in DIRS:
    out_dir = Path(DIRS["outputs"])
else:
    out_dir = Path("./outputs")
out_dir.mkdir(parents=True, exist_ok=True)

llm_dir = Path(DIRS["llm"]) if ("DIRS" in globals() and isinstance(DIRS, dict) and "llm" in DIRS) else Path("./llm")
llm_dir.mkdir(parents=True, exist_ok=True)

p_metrics = out_dir / "BACKTEST_METRICS.json"
p_exit = out_dir / "EXIT_EVENTS.csv"
p_llm_json = llm_dir / "llm_decision.json"

if not p_metrics.exists():
    raise FileNotFoundError(f"Missing: {p_metrics}. Run your backtest/metrics cell first.")

# ---------- Helpers ----------
def _to_float(x):
    try:
        return float(x)
    except Exception:
        return None

def _fmt(x, nd=2):
    v = _to_float(x)
    return f"{v:.{nd}f}" if v is not None else "N/A"

def _is_na(x):
    if x is None:
        return True
    s = str(x).strip().upper()
    return s in ["N/A", "NA", "NONE", ""]

def _annualized_sharpe(eq: pd.Series, rf: float = 0.0, periods: int = 252) -> float:
    r = eq.pct_change().dropna()
    if len(r) == 0 or r.std() == 0:
        return float("nan")
    return float(((r.mean() * periods) - rf) / (r.std() * np.sqrt(periods)))

def _max_drawdown_pct(eq: pd.Series) -> float:
    peak = eq.cummax()
    dd = (eq / peak - 1.0) * 100.0
    return float(dd.min())

def _cagr_pct(eq: pd.Series) -> float:
    eq = eq.dropna()
    if len(eq) < 2:
        return float("nan")
    years = (eq.index[-1] - eq.index[0]).days / 365.25
    if years <= 0:
        return float("nan")
    return float(((eq.iloc[-1] / eq.iloc[0]) ** (1.0 / years) - 1.0) * 100.0)

def _total_return_pct(eq: pd.Series) -> float:
    eq = eq.dropna()
    if len(eq) < 2:
        return float("nan")
    return float((eq.iloc[-1] / eq.iloc[0] - 1.0) * 100.0)

def _build_bh_equity(prices: pd.Series, start_equity: float) -> pd.Series:
    prices = prices.astype(float).ffill()
    shares = start_equity / float(prices.iloc[0])
    return shares * prices

def _safe_get(d, *keys, default=""):
    cur = d
    for k in keys:
        if not isinstance(cur, dict) or k not in cur:
            return default
        cur = cur[k]
    return cur if cur is not None else default

# ---------- Load metrics ----------
with open(p_metrics, "r", encoding="utf-8") as f:
    metrics = json.load(f)

s = metrics.get("strategy", {}) if isinstance(metrics, dict) else {}
asof = metrics.get("asof", str(df.index[-1].date()) if "df" in globals() else "N/A")

# backtest window (for possible CAGR backfill)
bt_start = None; bt_end = None
if "strategy_equity" in globals():
    try:
        bt_start = pd.to_datetime(strategy_equity.index[0])
        bt_end   = pd.to_datetime(strategy_equity.index[-1])
    except Exception:
        bt_start = bt_end = None
if bt_start is None or bt_end is None:
    try:
        bt_start = pd.to_datetime(s.get("bt_start"))
        bt_end   = pd.to_datetime(s.get("bt_end"))
    except Exception:
        pass
if (bt_start is None or bt_end is None) and "df" in globals():
    bt_start = pd.to_datetime(df.index[0])
    bt_end   = pd.to_datetime(df.index[-1])

bt_years = (bt_end - bt_start).days / 365.25 if (bt_start is not None and bt_end is not None) else None

# Strategy numbers
st_ret  = s.get("total_return_pct", None)
st_cagr = s.get("cagr_pct", None)
st_sh   = s.get("sharpe", None)
st_mdd  = s.get("max_drawdown_pct", None)
st_hit  = s.get("hit_rate_pct", None)
st_trd  = s.get("total_trades_closed", None)

# Backfill strategy CAGR if missing
if _is_na(st_cagr) and (not _is_na(st_ret)) and bt_years and bt_years > 0:
    total_mult = 1.0 + float(st_ret) / 100.0
    st_cagr = (total_mult ** (1.0 / bt_years) - 1.0) * 100.0

# ---------- Ensure B&H Adj Close exists (compute if missing) ----------
bh = metrics.get("buy_hold_adj_close", {}) if isinstance(metrics, dict) else {}
bh_ret  = bh.get("total_return_pct", None)
bh_cagr = bh.get("cagr_pct", None)
bh_sh   = bh.get("sharpe", None)
bh_mdd  = bh.get("max_drawdown_pct", None)

need_bh = any(_is_na(x) for x in [bh_ret, bh_cagr, bh_sh, bh_mdd])
if need_bh:
    if "df" not in globals() or "FAST_P" not in globals():
        raise NameError("Need df and FAST_P to compute Buy&Hold (Adj Close).")
    eval_df = df.iloc[int(FAST_P):].copy()
    start_equity = float(metrics.get("start_equity", globals().get("START_EQUITY", 10000)))
    rf = float(metrics.get("risk_free_rate", globals().get("RISK_FREE_RATE", 0.0)))
    eq_bh_adj = _build_bh_equity(eval_df["adj close"], start_equity=start_equity)
    bh_ret  = _total_return_pct(eq_bh_adj)
    bh_cagr = _cagr_pct(eq_bh_adj)
    bh_sh   = _annualized_sharpe(eq_bh_adj, rf=rf)
    bh_mdd  = _max_drawdown_pct(eq_bh_adj)

# ---------- Realised exits (show total + breakdown; minimal bullets) ----------
exit_lines = []
if p_exit.exists():
    try:
        ex = pd.read_csv(p_exit)
        total_exits = len(ex)
        col = None
        for c in ["reason", "exit_reason", "exit_type", "type"]:
            if c in ex.columns:
                col = c
                break
        exit_lines.append(f"- Total realised exits: {total_exits}")
        if col:
            vc = ex[col].astype(str).value_counts()
            for k, v in vc.head(4).items():
                exit_lines.append(f"- {k}: {int(v)}")
        else:
            exit_lines.append("- Breakdown not available (no reason/exit_reason/exit_type/type column found).")
    except Exception:
        exit_lines = ["- Total realised exits: N/A", "- Breakdown: parsing failed."]
else:
    exit_lines = ["- Total realised exits: N/A", "- EXIT_EVENTS.csv not found."]

# ---------- Latest snapshot ----------
ticker = globals().get("TICKER", "KO")
last = df.iloc[-1] if "df" in globals() else None
rec = str(last.get("recommendation", "HOLD")).upper() if last is not None else "HOLD"

# ---------- LLM JSON optional ----------
llm = {}
if p_llm_json.exists():
    try:
        llm = json.loads(p_llm_json.read_text(encoding="utf-8"))
    except Exception:
        llm = {}

abstract = _safe_get(llm, "Abstract", "Summary", default="").strip() or \
           (f"As of {asof}, the strategy signal for {ticker} is {rec}. "
            "The leakage-aware simulation (Close-executed) is benchmarked against Buy & Hold (Adj Close) as a dividend-inclusive proxy, "
            "and results are evaluated on risk-adjusted efficiency and drawdown control.")

rec_signal = _safe_get(llm, "Recommendation", "Signal", default="").strip() or rec.title()
rec_rat = _safe_get(llm, "Recommendation", "Rationale", default="").strip() or \
          ("The engine remains neutral because neither the full entry confluence nor the exit trigger set is satisfied at the latest bar, "
           "so the optimal action is to wait rather than force a trade.")

strategy_over = _safe_get(llm, "Strategy", "Overview", default="").strip() or \
                ("The strategy is a confluence-based, rule-driven technical system. It seeks mean-reversion entries when price recovers through "
                 "the Bollinger midline under constrained bias conditions, while requiring price to be below the fast moving average to avoid chasing. "
                 "Exits occur via TP, SL, or a signal-based bias exit; SL_FUSE provides intrabar protection against sudden adverse moves.")

ind_used = _safe_get(llm, "Strategy", "Indicators_used", default="").strip()
if not ind_used:
    ind_used = (
        f"fast_ma={_fmt(last.get('fast_ma'),4) if last is not None else 'N/A'}, "
        f"boll_mid/up/dn={_fmt(last.get('boll_mid'),4) if last is not None else 'N/A'}/"
        f"{_fmt(last.get('boll_up'),4) if last is not None else 'N/A'}/"
        f"{_fmt(last.get('boll_dn'),4) if last is not None else 'N/A'}, "
        f"bias={_fmt(last.get('bias'),4) if last is not None else 'N/A'}, "
        f"entry_signal={bool(last.get('entry_signal', False)) if last is not None else 'N/A'}, "
        f"exit_signal={bool(last.get('exit_signal', False)) if last is not None else 'N/A'}."
    )

bt_interp = _safe_get(llm, "Backtest_summary", "Performance_interpretation", default="").strip() or \
            ("Relative performance is assessed using both absolute returns and risk-adjusted efficiency. "
             "The key comparison emphasises Sharpe and drawdown to evaluate deployability under risk budgets.")

risk = _safe_get(llm, "Limitations_and_risks", "Discussion", default="").strip() or \
       ("Limitations include parameter sensitivity, regime dependence, and execution frictions (spreads, slippage, and gaps). "
        "Close-based execution is conservative for leakage control; Adj Close benchmarking embeds dividend effects that are not mechanically reinvested "
        "in the execution simulation, so interpretation should include this caveat.")

future = _safe_get(llm, "Future_refinements", "Discussion", default="").strip() or \
         ("Refinements should include ablation tests (remove TP, remove SL_FUSE, vary bias thresholds), walk-forward validation, "
          "and multi-asset / subperiod stress testing to verify stability under different regimes and transaction-cost assumptions.")

# ---------- Compose Markdown ----------
md = []
md.append(f"# {ticker} Trade Note ({asof})")
md.append("")
md.append("## Abstract")
md.append(abstract)
md.append("")
md.append("## Recommendation")
md.append(f"Signal: {rec_signal}")
md.append(f"Rationale: {rec_rat}")
md.append("")
md.append("## Strategy")
md.append(strategy_over)
md.append("")
md.append("## Indicators used")
md.append(ind_used)
md.append("")
md.append("## Backtest Summary")
md.append("Price-series disclosure: Strategy backtest/execution uses CLOSE (unadjusted) to avoid adjustment-related leakage; "
          "Buy & Hold uses ADJ CLOSE as a dividend-inclusive total-return proxy.")
md.append("")
md.append("- Strategy (Close): "
          f"Total Return {_fmt(st_ret,2)}%, CAGR {_fmt(st_cagr,2)}%, Sharpe {_fmt(st_sh,4)}, "
          f"Max Drawdown {_fmt(st_mdd,2)}%, Hit Rate {_fmt(st_hit,2)}%, Trades {st_trd if not _is_na(st_trd) else 'N/A'}")
md.append("- Buy & Hold (Adj Close): "
          f"Total Return {_fmt(bh_ret,2)}%, CAGR {_fmt(bh_cagr,2)}%, Sharpe {_fmt(bh_sh,4)}, "
          f"Max Drawdown {_fmt(bh_mdd,2)}%")
md.append("")
md.append(bt_interp)
md.append("")
md.append("### Realised exits")
md.extend(exit_lines)
md.append("")
md.append("## Limitations and Risks")
md.append(risk)
md.append("")
md.append("## Future Refinements")
md.append(future)
md.append("")

report_md = "\n".join(md).strip()

# ---------- Save MD ----------
p_report_md = out_dir / f"{REPORT_BASENAME}.md"
p_report_md.write_text(report_md, encoding="utf-8")
print("✅ Saved:", p_report_md)

# ---------- Save DOCX with larger titles + bigger gaps ----------
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

doc = Document()

sec = doc.sections[0]
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.75)
sec.right_margin = Inches(0.75)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)

def _apply_para(p, space_after_pt=10, line=1.05):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after_pt)

def _set_run_font(p, size_pt, bold=False):
    for r in p.runs:
        r.font.size = Pt(size_pt)
        r.bold = bold

def _blank_gap(lines=1):
    for _ in range(lines):
        p = doc.add_paragraph("")
        _apply_para(p, space_after_pt=0)

for line in report_md.splitlines():
    line = line.rstrip()
    if not line:
        continue

    if line.startswith("# "):
        p = doc.add_heading(line[2:], level=1)
        _apply_para(p, space_after_pt=18, line=1.0)
        _set_run_font(p, 18, bold=True)   # BIG title
        _blank_gap(1)

    elif line.startswith("## "):
        p = doc.add_heading(line[3:], level=2)
        _apply_para(p, space_after_pt=14, line=1.0)
        _set_run_font(p, 13, bold=True)   # Larger section title
        _blank_gap(1)

    elif line.startswith("### "):
        p = doc.add_heading(line[4:], level=3)
        _apply_para(p, space_after_pt=10, line=1.0)
        _set_run_font(p, 12, bold=True)
        _blank_gap(0)

    elif line.startswith("- "):
        p = doc.add_paragraph(line[2:], style="List Bullet")
        _apply_para(p, space_after_pt=4, line=1.05)  # keep bullets tight

    else:
        p = doc.add_paragraph(re.sub(r"\*\*(.*?)\*\*", r"\1", line))
        _apply_para(p, space_after_pt=10, line=1.05)

p_docx = out_dir / f"{REPORT_BASENAME}.docx"
doc.save(p_docx)
print("✅ Saved:", p_docx)

# In[ ]:




# In[ ]:




# In[ ]:



