# src/llm_trade_note.py
from __future__ import annotations

import os
from pathlib import Path
from typing import Any, Dict

from dotenv import load_dotenv
from docx import Document

try:
    from openai import AzureOpenAI
except Exception:
    AzureOpenAI = None  # type: ignore

load_dotenv()


def build_trade_note_prompt(result: Dict[str, Any], ticker: str) -> str:
    m = result["metrics"]
    df = result["df"].copy()

    last = df.iloc[-1]
    latest = {
        "date": str(df.index[-1].date()),
        "close": float(last["Close"]),
        "entry_cond": bool(last.get("entry_cond", False)),
        "buy_signal": bool(last.get("buy_signal", False)),
        "rsi50": float(last.get("rsi50", float("nan"))),
        "dist_pct": float(last.get("dist_pct", float("nan"))),
        "cmo15": float(last.get("cmo15", float("nan"))),
        "cmo_smma4": float(last.get("cmo_smma4", float("nan"))),
        "candle_ht_pct": float(last.get("candle_ht_pct", float("nan"))),
        "atr": float(last.get("atr", float("nan"))),
    }

    # IMPORTANT: coursework asks for hit rate; your metrics use win_rate.
    # We'll label it "hit rate" in the note but feed the value from win_rate.
    metrics_block = {
        "period": f'{m.get("start")} → {m.get("end")}',
        "CAGR": m.get("CAGR"),
        "Sharpe": m.get("Sharpe"),
        "MaxDrawdown": m.get("MaxDrawdown"),
        "HitRate(win_rate)": m.get("win_rate"),
        "num_trades": m.get("num_trades"),
        "profit_factor": m.get("profit_factor"),
        "avg_R": m.get("avg_R"),
        "observed_max_open_trades": m.get("observed_max_open_trades"),
        "observed_max_leverage": m.get("observed_max_leverage"),
        "risk_pct": m.get("risk_pct"),
        "atr_len": m.get("atr_len"),
        "sl_atr_mult": m.get("sl_atr_mult"),
        "rr": m.get("rr"),
        "slippage_bps": m.get("slippage_bps"),
    }

    return f"""
You are a buy-side Technical Analyst Agent in an asset management team.

Use ONLY the numbers provided (do not invent):
Backtest metrics: {metrics_block}
Latest snapshot: {latest}

ROUND ALL NUMBERS TO 2 DECIMAL PLACE

Write using exactly these headings:

TITLE: {ticker} — Technical Trade Note ({m.get("start")} to {m.get("end")})

Abstract
Max 100 words. State the analysis date and summarise recent performance using the latest indicators and signals.

Recommendation
Signal: Buy | Sell | Hold,
Rationale: Brief justification (1–2 sentences).

Strategy
Overview: Max 300 words explaining why the strategy generates this signal.


Backtest_summary
Performance_interpretation: Max 200 words interpreting results using key metrics (CAGR, Sharpe, Max Drawdown, Hit Rate, num_trades).

Current Signal & Trade Plan
Max 200 words. State whether entry_cond is true. If true, give a concrete plan:
- Entry timing (signal vs execution)
- Stop-loss and take-profit using ATR with sl_atr_mult and rr
- Position sizing using risk_pct of equity and how max_open_trades affects exposure.
If not true, state exactly what must change before entering.

Risks & Limitations
Around 100 words discussing weaknesses and risks.
""".strip()


def generate_trade_note_azure(result: Dict[str, Any], *, ticker: str) -> str:
    """
    Azure OpenAI trade note. No fallback.
    Requires env vars:
      AZURE_ENDPOINT
      AZURE_API_KEY
      AZURE_DEPLOYMENT_NAME
    Optional:
      AZURE_API_VERSION (default: 2024-02-15-preview)
    """
    if AzureOpenAI is None:
        raise RuntimeError("openai package with AzureOpenAI is not available. Install/upgrade openai.")

    # Fail fast if missing env vars
    missing = [k for k in ("AZURE_ENDPOINT", "AZURE_API_KEY", "AZURE_DEPLOYMENT_NAME") if not os.getenv(k)]
    if missing:
        raise RuntimeError(f"Missing required Azure env vars: {', '.join(missing)}")

    endpoint = os.environ["AZURE_ENDPOINT"]
    api_key = os.environ["AZURE_API_KEY"]
    deployment = os.environ["AZURE_DEPLOYMENT_NAME"]
    api_version = os.environ.get("AZURE_API_VERSION", "2024-02-15-preview")

    client = AzureOpenAI(
        azure_endpoint=endpoint,
        api_key=api_key,
        api_version=api_version,
    )

    prompt = build_trade_note_prompt(result, ticker)

    resp = client.chat.completions.create(
        model=deployment,  # Azure: deployment name
        messages=[
            {"role": "system", "content": "You write concise, data-grounded institutional research notes."},
            {"role": "user", "content": prompt},
        ],
        temperature=1,
    )

    content = resp.choices[0].message.content
    if not content:
        raise RuntimeError("Azure OpenAI returned empty trade note.")
    return content

# def save_trade_note_to_word(
#     trade_note: str,
#     *,
#     filename: str = "trade_note.docx",
#     output_dir: str = "output",
# ) -> Path:
#     """
#     Save trade note text to /output (same level as /src).
#     """
#     root = Path(__file__).resolve().parents[1]  # project root (../src -> root)
#     out_dir = root / output_dir
#     out_dir.mkdir(exist_ok=True)

#     path = out_dir / filename

#     doc = Document()
#     doc.add_heading("Trade Note", level=1)

#     for block in trade_note.split("\n\n"):
#         doc.add_paragraph(block)

#     doc.save(path)
#     return path
